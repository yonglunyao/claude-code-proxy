"""AIGC decorators for adding AIGC marks to existing documents."""
import hashlib
import json
import os.path
import stat
from struct import iter_unpack
from typing import Union, cast

import docx
from docx.shared import Pt
# Import docx_extend classes
from docx_extend.api import DocumentExtend
from openpyxl import load_workbook
from openpyxl.packaging.custom import StringProperty
from ppt_extend.aigc import add_aigc_mark_to_pptx
from pypdf import PdfReader, PdfWriter
from pypdf._utils import StreamType
from pypdf.generic import NameObject, TextStringObject, DictionaryObject

# Default fmtid for custom properties
DEFAULT_CUSTOM_PROPERTY_FMTID = "D5CDD505-2E9C-101B-9397-08002B2CF9AE"


def generate_sha256(text: str) -> str:
    """计算字符串的SHA256哈希值"""
    sha256_hash = hashlib.sha256()
    sha256_hash.update(text.encode('utf-8'))
    return sha256_hash.hexdigest()


def get_aigc_signature(content: str) -> str:
    """
    生成AIGC隐式标识

    Args:
        filename: 输入文件名（不含路径）

    Returns:
        JSON格式的AIGC标识字符串
    """
    # 计算文件名SHA256哈希的前16位作为内容ID
    content_id = generate_sha256(content)[:16]
    produce_id = f"voiceassistant-{content_id}"
    propagate_id = f"voiceassistant-{content_id}"

    aigc_metadata = {
        "Label": "1",
        "ContentProducer": "001191320114777023172010000",
        "ProduceID": produce_id,
        "ReservedCode1": "",
        "ContentPropagator": "001191320114777023172010000",
        "PropagateID": propagate_id
    }

    return json.dumps(aigc_metadata, ensure_ascii=False)


class RawTextStringObject(TextStringObject):
    """Custom TextStringObject that writes raw text without escaping."""

    def write_to_stream(
            self, stream: StreamType, encryption_key: Union[None, str, bytes] = None
    ) -> None:
        bytearr = self.get_encoded_bytes()
        stream.write(b"(")
        for c_ in iter_unpack("c", bytearr):
            c = cast(bytes, c_[0])
            stream.write(c)
        stream.write(b")")


class DocxAigcDecorator:
    """DOCX AIGC decorator - adds hidden AIGC mark to Word documents."""

    def __init__(self):
        self.name = "docx_aigc_decorator"

    def decorate(self, file_path: str, content: str, request_id: str = "", add_visible_mark: bool = True):
        """Add AIGC mark to DOCX file."""
        try:
            aigc_signature = get_aigc_signature(content)
            self._add_aigc_mark(file_path, aigc_signature, add_visible_mark)
            print(f"  [DOCX] AIGC mark added successfully")
        except Exception as e:
            print(f"  [DOCX] Warning: Failed to add AIGC mark: {str(e)}")

    def _add_aigc_mark(self, file_path: str, signature: str, add_visible_mark: bool = True):
        """Add custom property to DOCX file using docx library."""
        if not file_path or not os.path.exists(file_path):
            raise FileNotFoundError(f"Target file not found: {file_path}")

        doc = docx.Document(file_path)
        if add_visible_mark:
            self._add_visible_mark_to_docx(doc)
        self._add_implicit_mark(doc, signature)
        doc.save(file_path)

    def _add_visible_mark_to_docx(self, doc: docx.Document):
        """Add visible AIGC mark paragraph to DOCX file."""

        # 添加段落
        paragraph = doc.add_paragraph()
        run = paragraph.add_run("内容由AI生成")

        # 设置字体为宋体，5号字(10.5pt)
        run.font.name = "宋体"
        run.font.size = Pt(10.5)

    def _add_implicit_mark(self, doc, signature):
        doc_ex = DocumentExtend(doc)
        custom_properties_part = doc_ex.custom_properties_part
        pid = custom_properties_part.next_id
        custom_properties = custom_properties_part.custom_properties
        custom_properties.add_property("AIGC", signature, DEFAULT_CUSTOM_PROPERTY_FMTID, pid)


class PdfAigcDecorator:
    """PDF AIGC decorator - adds hidden AIGC mark to PDF documents."""

    def __init__(self):
        self.name = "pdf_aigc_decorator"

    def decorate(self, file_path: str, content: str, request_id: str = "", add_visible_mark: bool = True):
        """Add AIGC mark to PDF file."""
        try:
            aigc_signature = get_aigc_signature(content)
            aigc_signature = aigc_signature.replace("\\", "\\\\")
            self._add_aigc_mark(file_path, {"AIGC": aigc_signature, "Creator": "", "Producer": ""}, add_visible_mark)
            print(f"  [PDF] AIGC mark added successfully")
        except Exception as e:
            print(f"  [PDF] Warning: Failed to add AIGC mark: {str(e)}")

    def _add_aigc_mark(self, input_path: str, data: dict, add_visible_mark: bool = True) -> None:
        """Add or update metadata and optional watermark in a PDF file."""
        # Read input PDF
        reader = PdfReader(input_path)
        writer = PdfWriter()

        # Copy all pages
        writer.clone_reader_document_root(reader)

        # Add visible watermark if requested
        if add_visible_mark:
            self._add_visible_watermark(writer, reader)

        # Ensure writer._info is a DictionaryObject
        if writer._info is None:
            writer._info = DictionaryObject()

        # Preserve existing metadata
        if reader.metadata:
            for key, value in reader.metadata.items():
                pdf_key = NameObject(key)
                if pdf_key not in writer._info:
                    writer._info[pdf_key] = RawTextStringObject(str(value))

        # Add/update new metadata
        for key, value in data.items():
            pdf_key = NameObject(f"/{key.lstrip('/')}")
            writer._info[pdf_key] = RawTextStringObject(str(value))

        # Write output (overwrite input file)
        file_flags = os.O_WRONLY | os.O_CREAT
        file_mode = stat.S_IWUSR | stat.S_IRUSR
        try:
            with os.fdopen(os.open(input_path, file_flags, file_mode), 'wb') as f:
                writer.write(f)
        except Exception as e:
            raise Exception(f"PDF add metadata failed: {str(e)}") from e

    def _add_visible_watermark(self, writer: PdfWriter, reader: PdfReader):
        """Add visible '内容由AI生成' watermark to PDF pages using reportlab-generated watermark."""
        import os
        from pypdf import PdfReader as LocalPdfReader

        # 水印文件路径
        watermark_path = os.path.join(
            os.path.dirname(__file__), '..', 'pdf_extend', 'aigc_watermark.pdf'
        )
        watermark_path = os.path.abspath(watermark_path)

        # 如果水印文件不存在，生成它
        if not os.path.exists(watermark_path):
            from pdf_extend.pdf_watermark import create_watermark_pdf
            create_watermark_pdf(
                output_path=watermark_path,
                text="内容由AI生成",
                font_size=10.5,
                opacity=1.0,
                angle=0,
                color=(0, 0, 0),  # 黑色
                position='bottom-center'
            )

        # 读取水印
        watermark_reader = LocalPdfReader(watermark_path)
        watermark_page = watermark_reader.pages[0]

        # 为每一页添加水印（合并到底层）
        for page_num in range(len(reader.pages)):
            page = writer.pages[page_num]
            # 将水印页合并到当前页（水印在底层）
            page.merge_page(watermark_page, over=False)


class ExcelAigcDecorator:
    """Excel AIGC decorator - adds hidden AIGC mark to Excel files."""

    def __init__(self):
        self.name = "excel_aigc_decorator"

    def decorate(self, file_path: str, content: str, request_id: str = "", add_visible_mark: bool = True):
        """Add AIGC mark to Excel file."""
        try:
            aigc_signature = get_aigc_signature(content)
            self._add_aigc_mark(file_path, aigc_signature, add_visible_mark)
            print(f"  [Excel] AIGC mark added successfully")
        except Exception as e:
            print(f"  [Excel] Warning: Failed to add AIGC mark: {str(e)}")

    def _add_aigc_mark(self, file_path: str, signature: str, add_visible_mark: bool = True):
        """Add custom property to Excel file."""
        wb = load_workbook(file_path)
        if add_visible_mark:
            self.add_visible_mark(wb)
        self._add_implicit_mark(wb, signature)
        wb.save(file_path)

    def _add_implicit_mark(self, wb, signature):
        custom_props = wb.custom_doc_props
        if custom_props is not None:
            custom_props.append(StringProperty(name='AIGC', value=signature))

    def add_visible_mark(self, workbook):
        """
        添加显示AI标识

        在每个sheet的最后一行添加"内容由AI生成"
        """
        ai_mark_content = "内容由AI生成"

        for sheet_name in workbook.sheetnames:
            ws = workbook[sheet_name]
            ws.append([ai_mark_content])


class PptAigcDecorator:
    """PPT AIGC decorator - adds hidden AIGC mark to PowerPoint presentations."""

    def __init__(self):
        self.name = "ppt_aigc_decorator"

    def decorate(self, file_path: str, content: str, request_id: str = "", add_visible_mark: bool = True):
        """Add AIGC mark to PPTX file."""
        try:
            aigc_signature = get_aigc_signature(content)
            add_aigc_mark_to_pptx(file_path, file_path, aigc_signature, request_id, add_visible_mark)
            print(f"  [PPT] AIGC mark added successfully")
        except Exception as e:
            print(f"  [PPT] Warning: Failed to add AIGC mark: {str(e)}")


class MdAigcDecorator:
    """MD AIGC decorator - adds AIGC mark to Markdown files."""

    def __init__(self):
        self.name = "md_aigc_decorator"

    def decorate(self, file_path: str, content: str, request_id: str = "", add_visible_mark: bool = True):
        """Add AIGC mark to Markdown file."""
        try:
            aigc_signature = get_aigc_signature(content)
            metadata = json.loads(aigc_signature)
            self._add_aigc_mark(file_path, metadata, add_visible_mark)
            print(f"  [MD] AIGC mark added successfully")
        except Exception as e:
            print(f"  [MD] Warning: Failed to add AIGC mark: {str(e)}")

    def _add_aigc_mark(self, file_path, metadata: dict, add_visible_mark: bool = True):
        try:
            markdown_content = ""
            with open(file_path, "r", encoding="utf-8") as f:
                markdown_content = f.read()
            label = metadata.get("Label", "")
            content_producer = metadata.get("ContentProducer", "")
            produce_id = metadata.get("ProduceID", "")
            reserved_code1 = metadata.get("ReservedCode1", "")
            content_propagator = metadata.get("ContentPropagator", "")
            propagator_id = metadata.get("PropagateID", "")
            aigc_signature = (f"---\nAIGC:\n"
                              f"  Label: {label}\n"
                              f"  ContentProducer: {content_producer}\n"
                              f"  ProduceID: {produce_id}\n"
                              f"  ReservedCode1: {reserved_code1}\n"
                              f"  ContentPropagator: {content_propagator}\n"
                              f"  PropagateID: {propagator_id}\n---")

            # Build content based on add_visible_mark
            if add_visible_mark:
                new_content = f"{aigc_signature}\n\n{markdown_content}\n\n内容由AI生成"
            else:
                new_content = f"{aigc_signature}\n\n{markdown_content}"

            file_flags = os.O_WRONLY | os.O_CREAT
            file_mode = stat.S_IWUSR | stat.S_IRUSR
            with os.fdopen(os.open(file_path, file_flags, file_mode), 'w', encoding='utf-8') as f:
                f.write(new_content)
        except Exception as e:
            print(f' [MD] add aigc mark error due to: {str(e)}')
            return markdown_content


# Create singleton instances
docx_aigc_decorator = DocxAigcDecorator()
pdf_aigc_decorator = PdfAigcDecorator()
excel_aigc_decorator = ExcelAigcDecorator()
ppt_aigc_decorator = PptAigcDecorator()
md_aigc_decorator = MdAigcDecorator()
